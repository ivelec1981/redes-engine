# -*- coding: utf-8 -*-
"""
Tests del módulo reports — generación PDF y Word.
"""

import os
import sys
from datetime import datetime

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)

import pytest

# Skips si las libs no están
reportlab = pytest.importorskip("reportlab", reason="reportlab no instalado")
docx = pytest.importorskip("docx", reason="python-docx no instalado")
mpl = pytest.importorskip("matplotlib", reason="matplotlib no instalado")


from redes_engine.reports import (
    ReportBuilder, ReportContext, generate_docx_report, generate_pdf_report,
)
from redes_engine.reports import charts as chart_module


# =============================================================================
# ReportContext
# =============================================================================
class TestReportContext:

    def test_empty_context(self):
        ctx = ReportContext(title="X")
        assert ctx.title == "X"
        assert ctx.has_solve is False
        assert ctx.has_compliance is False

    def test_context_flags(self):
        ctx = ReportContext(title="X", flow_result="dummy")
        assert ctx.has_solve is True


# =============================================================================
# Builder
# =============================================================================
class TestReportBuilder:

    @pytest.fixture
    def basic_ctx(self):
        from redes_engine.examples.urbanizacion_mixta import (
            build_urbanizacion_pastaza,
        )
        return ReportContext(
            title="Test", project_name="P1",
            network=build_urbanizacion_pastaza(),
        )

    def test_builds_summary_section(self, basic_ctx):
        sections = ReportBuilder(basic_ctx).build_all()
        # Al menos resumen + descripción de la red + recomendaciones
        titles = [s.title for s in sections]
        assert any("Resumen" in t for t in titles)
        assert any("Descripción" in t for t in titles)

    def test_skip_section_without_data(self):
        ctx = ReportContext(title="Empty")
        sections = ReportBuilder(ctx).build_all()
        # Sin red, sin solve, sin compliance → solo Resumen + Recomendaciones
        titles = [s.title for s in sections]
        assert "Resumen Ejecutivo" in titles
        assert not any("Flujo de Potencia" in t for t in titles)

    def test_investment_section_appears_when_provided(self):
        from redes_engine.engineering import (
            InvestmentAnalyzer, InvestmentAssumptions,
        )
        ana = InvestmentAnalyzer(InvestmentAssumptions(horizon_years=5))
        inv = ana.analyze(
            capex_direct_usd=180_000,
            annual_loss_savings_kwh=40_000,
            annual_capacity_savings_usd=18_000,
        )
        ctx = ReportContext(
            title="Inv", project_name="X",
            investment_result=inv,
        )
        assert ctx.has_investment is True
        sections = ReportBuilder(ctx).build_all()
        titles = [s.title for s in sections]
        assert any("Inversión" in t for t in titles)
        # Verificar que la sección tiene métricas
        inv_sec = next(s for s in sections if "Inversión" in s.title)
        # Indicadores y flujos = 2 tablas
        assert len(inv_sec.tables) == 2

    def test_bom_section_appears_when_provided(self):
        from redes_engine.engineering import BudgetEngine, UPItem
        engine = BudgetEngine(
            uc_database={"PSC": {"materials": [
                {"item": "X1", "descripcion": "X", "unidad": "u",
                 "cantidad": 1},
            ]}},
            prices_database={"X1": 50.0},
        )
        bom = engine.compute_bom([UPItem(code="PSC", quantity=2)])
        ctx = ReportContext(title="B", bom_result=bom)
        assert ctx.has_bom is True
        sections = ReportBuilder(ctx).build_all()
        assert any("Presupuesto" in s.title for s in sections)


# =============================================================================
# Charts
# =============================================================================
class TestCharts:

    def test_voltage_chart_returns_png_bytes(self):
        from redes_engine.core.results import (
            BusVoltageResult, ComplianceStatus, PowerFlowResult,
        )
        r = PowerFlowResult(converged=True)
        r.bus_voltages["B1"] = BusVoltageResult(
            bus_id="B1", v_magnitude_kv=22.5, v_pu=0.987,
            v_drop_pct=1.3, angle_deg=0.0, voltage_nominal_kv=22.8,
            compliance=ComplianceStatus.OK,
        )
        r.bus_voltages["B2"] = BusVoltageResult(
            bus_id="B2", v_magnitude_kv=20.5, v_pu=0.90,
            v_drop_pct=10.0, angle_deg=0.0, voltage_nominal_kv=22.8,
            compliance=ComplianceStatus.VIOLATION,
        )
        png = chart_module.bar_voltage_profile(r)
        assert isinstance(png, bytes)
        # PNG magic bytes
        assert png[:8] == b"\x89PNG\r\n\x1a\n"

    def test_branch_loading_chart(self):
        from redes_engine.core.results import (
            BranchFlowResult, ComplianceStatus, PowerFlowResult,
        )
        r = PowerFlowResult(converged=True)
        r.branch_flows["L1"] = BranchFlowResult(
            branch_id="L1", p_kw=10, q_kvar=2, s_kva=10,
            current_a=50, rated_a=100, loading_pct=50,
            losses_kw=0.1, losses_kvar=0.0,
            compliance=ComplianceStatus.OK,
        )
        png = chart_module.bar_branch_loading(r)
        assert png[:8] == b"\x89PNG\r\n\x1a\n"


# =============================================================================
# Generación PDF
# =============================================================================
class TestPDFGeneration:

    def test_pdf_with_minimal_context(self, tmp_path):
        ctx = ReportContext(
            title="Mínimo",
            project_name="Test",
            issue_date=datetime(2026, 5, 6),
        )
        path = tmp_path / "min.pdf"
        result_path = generate_pdf_report(ctx, str(path))
        assert os.path.exists(result_path)
        assert os.path.getsize(result_path) > 1000   # PDF razonable

    def test_pdf_with_network(self, tmp_path):
        from redes_engine.examples.urbanizacion_mixta import (
            build_urbanizacion_pastaza,
        )
        ctx = ReportContext(
            title="Red Pastaza",
            project_name="P1",
            network=build_urbanizacion_pastaza(),
            include_charts=False,   # sin charts para test rápido
        )
        path = tmp_path / "red.pdf"
        generate_pdf_report(ctx, str(path))
        assert os.path.getsize(path) > 1000

    def test_pdf_full_pipeline(self, tmp_path):
        """PDF con red + solve + compliance + hosting + annual."""
        from redes_engine.examples.urbanizacion_mixta import (
            build_urbanizacion_pastaza,
        )
        from redes_engine.io.opendss_solver import OpenDSSSolver
        from redes_engine.core.compliance import (
            ARCERNNR_EC, ComplianceAnalyzer,
        )

        net = build_urbanizacion_pastaza()
        with OpenDSSSolver(net) as solver:
            solver.solve()
            flow = solver.collect_results()
        compliance = ComplianceAnalyzer(ARCERNNR_EC).analyze(flow)

        ctx = ReportContext(
            title="Pipeline completo", project_name="Pastaza",
            network=net, flow_result=flow,
            compliance_report=compliance,
            include_charts=True,
        )
        path = tmp_path / "full.pdf"
        generate_pdf_report(ctx, str(path))
        # Con charts y datos completos debe ser un PDF razonable
        assert os.path.getsize(path) > 30000


# =============================================================================
# Generación DOCX
# =============================================================================
class TestDocxGeneration:

    def test_docx_minimal(self, tmp_path):
        ctx = ReportContext(title="Mínimo Word")
        path = tmp_path / "min.docx"
        generate_docx_report(ctx, str(path))
        assert os.path.exists(path)
        assert os.path.getsize(path) > 5000

    def test_docx_with_data(self, tmp_path):
        from redes_engine.examples.urbanizacion_mixta import (
            build_urbanizacion_pastaza,
        )
        ctx = ReportContext(
            title="Test Word", project_name="P1",
            network=build_urbanizacion_pastaza(),
            include_charts=False,
        )
        path = tmp_path / "data.docx"
        generate_docx_report(ctx, str(path))
        # Verificar que es un docx válido (lo abrimos)
        from docx import Document
        d = Document(str(path))
        # Debe haber secciones
        all_text = "\n".join(p.text for p in d.paragraphs)
        assert "Test Word" in all_text or "P1" in all_text


# =============================================================================
# Endpoint API
# =============================================================================
fastapi_pkg = pytest.importorskip("fastapi", reason="fastapi no instalado")
opendss = pytest.importorskip("opendssdirect", reason="opendssdirect no instalado")


class TestReportEndpoint:

    @pytest.fixture(autouse=True)
    def clean_store(self):
        from redes_engine.api.storage import get_store
        get_store().clear()
        yield
        get_store().clear()

    @pytest.fixture
    def client(self):
        from fastapi.testclient import TestClient
        from redes_engine.api.main import app
        return TestClient(app)

    def test_report_endpoint_pdf(self, client):
        # Cargar red + solve + report
        nid = client.post("/api/v1/demo/load").json()["id"]
        client.post(f"/api/v1/networks/{nid}/solve")
        res = client.post(
            f"/api/v1/networks/{nid}/report",
            json={"format": "pdf", "project_name": "Test PDF"},
        )
        assert res.status_code == 200
        assert res.headers["content-type"] == "application/pdf"
        assert len(res.content) > 5000

    def test_report_endpoint_docx(self, client):
        nid = client.post("/api/v1/demo/load").json()["id"]
        client.post(f"/api/v1/networks/{nid}/solve")
        res = client.post(
            f"/api/v1/networks/{nid}/report",
            json={"format": "docx"},
        )
        assert res.status_code == 200
        assert "wordprocessingml" in res.headers["content-type"]

    def test_report_without_solve_400(self, client):
        nid = client.post("/api/v1/demo/load").json()["id"]
        # Sin hacer solve antes
        res = client.post(
            f"/api/v1/networks/{nid}/report",
            json={"format": "pdf"},
        )
        assert res.status_code == 400
