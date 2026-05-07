# -*- coding: utf-8 -*-
"""
redes_engine.reports.executive_docx
=====================================

Generador de reporte ejecutivo en formato Word (.docx) usando python-docx.

A diferencia del PDF (firmable), el Word permite edición posterior por el
ingeniero responsable o por el cliente.
"""

import io
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from . import charts
from .builder import ReportBuilder, ReportContext, ReportSection

# =============================================================================
# Helpers de estilo
# =============================================================================
COLOR_TITLE   = RGBColor(0x1A, 0x23, 0x7E)
COLOR_HEADING = RGBColor(0x1A, 0x23, 0x7E)
COLOR_GRAY    = RGBColor(0x5A, 0x63, 0x73)


def _set_cell_bg(cell, hex_color: str):
    """Color de fondo de una celda (hex sin '#')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_heading(doc, text: str, level: int = 1):
    """Encabezado con color personalizado."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADING
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _add_body(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


# =============================================================================
# Portada
# =============================================================================
def _add_cover(doc, ctx: ReportContext):
    # Espacio superior
    for _ in range(3):
        doc.add_paragraph()

    # Logo si existe
    if ctx.company_logo:
        try:
            doc.add_picture(ctx.company_logo, width=Cm(4))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ctx.title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_TITLE

    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ctx.subtitle)
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    # Tabla de metadatos
    rows = [
        ("Proyecto", ctx.project_name or "—"),
        ("Empresa", ctx.company_name),
        ("Responsable", ctx.author_name),
        ("Licencia", ctx.author_id or "—"),
        ("Email", ctx.author_email or "—"),
        ("Código documento", ctx.document_code or "—"),
        ("Revisión", ctx.revision),
        ("Fecha de emisión", ctx.issue_date.strftime("%d/%m/%Y")),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False

    for i, (k, v) in enumerate(rows):
        c1 = table.cell(i, 0)
        c2 = table.cell(i, 1)
        c1.width = Cm(5)
        c2.width = Cm(9)
        _set_cell_bg(c1, "E3F2FD")
        for run_p in c1.paragraphs:
            r = run_p.add_run(k)
            r.font.bold = True
            r.font.color.rgb = COLOR_HEADING
            r.font.size = Pt(10)
        for run_p in c2.paragraphs:
            r = run_p.add_run(str(v))
            r.font.size = Pt(10)

    # Línea de cierre
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Generado por redes_engine — Análisis ARCERNNR Reg. 002/20"
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = COLOR_GRAY

    doc.add_page_break()


# =============================================================================
# Render de una sección
# =============================================================================
def _add_section(doc, sec: ReportSection, ctx: ReportContext):
    _add_heading(doc, sec.title, level=1)

    for para in sec.body_paragraphs:
        _add_body(doc, para)

    for tbl in sec.tables:
        if tbl.get("title"):
            p = doc.add_paragraph()
            r = p.add_run(tbl["title"])
            r.font.bold = True
            r.font.size = Pt(11)

        n_cols = len(tbl["headers"])
        n_rows = 1 + len(tbl["rows"])
        word_tbl = doc.add_table(rows=n_rows, cols=n_cols)
        word_tbl.style = "Light Grid Accent 1"
        # Header
        hdr_cells = word_tbl.rows[0].cells
        for i, h in enumerate(tbl["headers"]):
            cell = hdr_cells[i]
            _set_cell_bg(cell, "1A237E")
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(str(h))
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(9)
        # Filas
        for ridx, row in enumerate(tbl["rows"], start=1):
            for cidx, val in enumerate(row):
                cell = word_tbl.rows[ridx].cells[cidx]
                cell.paragraphs[0].clear()
                r = cell.paragraphs[0].add_run(str(val))
                r.font.size = Pt(9)

    # Charts
    if ctx.include_charts:
        _add_chart_for_section(doc, sec.title, ctx)


def _add_chart_for_section(doc, sec_title: str, ctx: ReportContext):
    if "Flujo de Potencia" in sec_title and ctx.has_solve:
        try:
            png1 = charts.bar_voltage_profile(ctx.flow_result)
            png2 = charts.bar_branch_loading(ctx.flow_result)
            doc.add_picture(io.BytesIO(png1), width=Cm(16))
            doc.add_picture(io.BytesIO(png2), width=Cm(16))
        except Exception:
            pass
    elif "Capacidad de Alojamiento" in sec_title and ctx.has_hosting:
        try:
            png = charts.hosting_capacity_chart(ctx.hosting_results)
            doc.add_picture(io.BytesIO(png), width=Cm(16))
        except Exception:
            pass


# =============================================================================
# ENTRADA PÚBLICA
# =============================================================================
def generate_docx_report(
    context: ReportContext,
    output_path: str,
) -> str:
    """
    Genera un reporte Word con todas las secciones disponibles.

    Returns
    -------
    str : ruta al .docx generado.
    """
    doc = Document()
    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    _add_cover(doc, context)

    builder = ReportBuilder(context)
    for sec in builder.build_all():
        _add_section(doc, sec, context)

    if context.extra_notes:
        _add_heading(doc, "Notas adicionales", level=1)
        _add_body(doc, context.extra_notes)

    doc.save(output_path)
    return output_path
